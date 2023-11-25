import boto3
from docx import Document
import sys


def read_docx(file_path):
    doc = Document(file_path)
    full_text = []
    for paragraph in doc.paragraphs:
        full_text.append(paragraph.text)
    return ' '.join(full_text)


def translate_text_block(block, translate_client):
    return translate_client.translate_text(Text=block, SourceLanguageCode='en', TargetLanguageCode='es')[
        'TranslatedText']


def write_translated_docx(text, original_file_path):
    new_file_path = original_file_path.rsplit('.', 1)[0] + '_spanish.docx'
    doc = Document()
    doc.add_paragraph(text)
    doc.save(new_file_path)
    return new_file_path


def convert_docx_to_translated_docx():
    file_path = input('Please enter the path to the Word document or type "exit" to abort: ')

    if file_path.lower() == 'exit':
        print('Program aborted by the user.')
        sys.exit()

    text = read_docx(file_path)

    # Dividir el texto en bloques de 4000 caracteres
    max_block_len = 4000
    blocks = [text[i:i + max_block_len] for i in range(0, len(text), max_block_len)]

    # Crear el cliente de AWS Translate
    translate_client = boto3.client(service_name='translate', region_name='eu-central-1')

    # Traducir cada bloque y acumular los resultados
    translated_blocks = ''
    for i, block in enumerate(blocks, start=1):
        translated_block = translate_text_block(block, translate_client)
        translated_blocks += translated_block

        # Mostrar el progreso
        sys.stdout.write(f'\rTranslating block {i}/{len(blocks)}...')
        sys.stdout.flush()

    sys.stdout.write('\n')

    new_file_path = write_translated_docx(translated_blocks, file_path)

    print(f'Translation completed. Translated document saved as {new_file_path}')


convert_docx_to_translated_docx()
